from docx import Document

# Create the document
doc = Document()
doc.add_heading("Supporting Statement", level=1)

# Applicant Info
doc.add_paragraph("Applicant: Peter Kibara Muriuki")
doc.add_paragraph("Position: Investigations Team Internship (Internship 1) – Office of Internal Audit and Investigations (OIAI)")
doc.add_paragraph("Duty Station: Nairobi (or Remote)")
doc.add_paragraph("")

# Motivation and Interest
doc.add_heading("Motivation and Interest", level=2)
doc.add_paragraph(
    "I am motivated to join UNICEF’s Office of Internal Audit and Investigations because I believe integrity and transparency are fundamental "
    "to achieving social impact. I am particularly drawn to OIAI’s role in ensuring that UNICEF operations uphold the highest standards of "
    "ethical conduct, accountability, and safeguarding.\n\n"
    "My academic background in Business Information Technology has strengthened my capacity to conduct structured analysis, manage data ethically, "
    "and apply technology to support oversight and compliance functions. I am eager to contribute my analytical and research skills to help "
    "identify trends, verify data accuracy, and support the investigative process."
)

# Relevant Skills and Experience
doc.add_heading("Relevant Skills and Experience", level=2)
skills = [
    "Analytical Research: Conducted structured research and data analysis projects involving data verification, report preparation, and risk identification.",
    "Information Management: Experienced in handling confidential information and maintaining data integrity in project and client systems.",
    "Digital Competence: Skilled in Microsoft Office Suite, Excel (Advanced), and open-source research methods for data verification.",
    "Communication: Strong technical writing and documentation skills; able to produce clear and concise reports and summaries for multiple audiences.",
    "Ethical Awareness: Adhere strictly to confidentiality, impartiality, and integrity in handling sensitive data and information.",
]
for s in skills:
    doc.add_paragraph(s, style="List Bullet")

# Alignment with UNICEF Competencies
doc.add_heading("Alignment with UNICEF Competencies", level=2)
competencies = [
    "Integrity & Ethical Awareness: I apply ethical judgment and respect for confidentiality in all professional activities.",
    "Analytical Thinking: I use structured approaches to identify, analyze, and solve complex information-related problems.",
    "Accountability & Results Orientation: I take ownership of assigned tasks, meet deadlines, and maintain accuracy in all deliverables.",
    "Teamwork: I work effectively in multicultural and multidisciplinary environments, contributing positively to team objectives.",
    "Adaptability: I can manage multiple tasks simultaneously and respond constructively to evolving priorities or new procedures.",
]
for c in competencies:
    doc.add_paragraph(c, style="List Bullet")

# Core Values
doc.add_heading("Core Values (CRITAS)", level=2)
doc.add_paragraph(
    "I am committed to UNICEF’s Core Values of Care, Respect, Integrity, Trust, Accountability, and Sustainability. "
    "These values guide how I approach my work — with diligence, transparency, and a sense of responsibility toward the communities UNICEF serves."
)

# Availability and Commitment
doc.add_heading("Availability and Commitment", level=2)
doc.add_paragraph(
    "I am available to begin the internship immediately and can work in-person in Nairobi or remotely as required. "
    "I am committed to upholding UNICEF’s zero-tolerance policy on misconduct and contributing positively to its mission of protecting and empowering children globally."
)

# Declaration
doc.add_heading("Declaration", level=2)
doc.add_paragraph(
    "I certify that all information provided is accurate to the best of my knowledge and that I am committed to maintaining the highest standards "
    "of ethics, professionalism, and confidentiality during my internship with UNICEF."
)

doc.add_paragraph("\nSignature:\tPeter Kibara Muriuki")
doc.add_paragraph("Date:\t18 October 2025")

# Save the file
file_path = "/mnt/data/Peter_Kibara_UNICEF_Supporting_Statement.docx"
doc.save(file_path)

file_path
